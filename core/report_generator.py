"""
core/report_generator.py
========================
Genera el pre-informe de liquidación como documento DOCX (Word) usando python-docx.
El usuario puede abrirlo en Word/LibreOffice y editarlo libremente.
"""

from __future__ import annotations

import io

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# Constantes de estilo
# ---------------------------------------------------------------------------
AZUL = RGBColor(0x1A, 0x3A, 0x6B)
ROJO = RGBColor(0xB1, 0x00, 0x00)
VERDE = RGBColor(0x0A, 0x52, 0x00)
NARANJA = RGBColor(0xB4, 0x5A, 0x00)
GRIS = RGBColor(0x55, 0x55, 0x55)


# ---------------------------------------------------------------------------
# Helpers internos
# ---------------------------------------------------------------------------

def _shading(cell, fill_hex: str) -> None:
    """Aplica color de fondo a una celda de tabla."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _cell_border_bottom(cell) -> None:
    """Agrega borde inferior a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "1A3A6B")
    tcBorders.append(bottom)
    tcPr.append(tcBorders)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11 if level == 1 else 10)
    run.font.color.rgb = AZUL
    run.underline = level == 1


def _add_data_row(table, label: str, value: str) -> None:
    """Agrega una fila etiqueta/valor a una tabla de dos columnas."""
    row = table.add_row()
    row.height = Pt(18)

    cell_lbl = row.cells[0]
    _shading(cell_lbl, "F0F4FA")
    p_lbl = cell_lbl.paragraphs[0]
    p_lbl.paragraph_format.space_before = Pt(2)
    p_lbl.paragraph_format.space_after = Pt(2)
    run = p_lbl.add_run(label)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9)

    cell_val = row.cells[1]
    p_val = cell_val.paragraphs[0]
    p_val.paragraph_format.space_before = Pt(2)
    p_val.paragraph_format.space_after = Pt(2)
    run2 = p_val.add_run(str(value) if value else "—")
    run2.font.name = "Arial"
    run2.font.size = Pt(9)


def _bullet(doc: Document, text: str, color=None) -> None:
    """Agrega un párrafo con viñeta."""
    try:
        p = doc.add_paragraph(style="List Bullet")
    except KeyError:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run("• ")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = color


def _fmt_vehiculo(v: dict) -> str:
    if not v or not v.get("marca"):
        return "—"
    parts = [f'{v.get("marca", "")} {v.get("modelo", "")}'.strip()]
    if v.get("año"):
        parts.append(f'año {v["año"]}')
    if v.get("color"):
        parts.append(f'color {v["color"]}')
    if v.get("matricula"):
        parts.append(f'Chapa N° {v["matricula"]}')
    if v.get("vin"):
        parts.append(f'VIN N° {v["vin"]}')
    return ", ".join(filter(None, parts))


def _fmt_conductor(c: dict) -> str:
    if not c or not c.get("nombre"):
        return "—"
    s = c["nombre"]
    if c.get("ci"):
        s += f' C.I. N° {c["ci"]}'
    if c.get("licencia"):
        s += f', Lic. N° {c["licencia"]}'
    return s


# ---------------------------------------------------------------------------
# Función principal
# ---------------------------------------------------------------------------

def generar_docx(datos: dict) -> bytes:
    """
    Recibe el dict de `generar_datos_preinforme()` y devuelve los bytes del DOCX.

    Usage:
        docx_bytes = generar_docx(datos)
        HttpResponse(docx_bytes, content_type='application/vnd.openxmlformats-...')
    """
    doc = Document()

    # ── Página A4 ──────────────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)

    # ── Fuente por defecto ─────────────────────────────────────────────────
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Arial"
    style_normal.font.size = Pt(10)

    caso_id = datos.get("caso_id", "")

    # ====================================================================
    # CABECERA
    # ====================================================================
    p_cab = doc.add_paragraph()
    p_cab.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_cab.add_run("SISTEMA DE LIQUIDACIÓN DE SINIESTROS — INFORME PRELIMINAR")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = AZUL

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run(
        f'Caso: {caso_id}   |   Generado: {datos.get("generado_en", "")}'
    )
    r2.font.size = Pt(8)
    r2.font.color.rgb = GRIS
    p_sub.paragraph_format.space_after = Pt(6)

    # ── Línea separadora (párrafo con borde inferior) ─────────────────────
    p_sep = doc.add_paragraph()
    pPr = p_sep._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom_bdr = OxmlElement("w:bottom")
    bottom_bdr.set(qn("w:val"), "single")
    bottom_bdr.set(qn("w:sz"), "12")
    bottom_bdr.set(qn("w:color"), "1A3A6B")
    pBdr.append(bottom_bdr)
    pPr.append(pBdr)
    p_sep.paragraph_format.space_after = Pt(8)

    # ====================================================================
    # TÍTULO PRINCIPAL
    # ====================================================================
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_titulo = p_titulo.add_run("PRE-INFORME")
    r_titulo.bold = True
    r_titulo.font.size = Pt(16)
    r_titulo.font.color.rgb = AZUL
    r_titulo.underline = True
    p_titulo.paragraph_format.space_after = Pt(14)

    # ====================================================================
    # TABLA PRINCIPAL DE DATOS
    # ====================================================================
    _heading(doc, "Datos del Siniestro")

    veh_a = datos.get("vehiculo_asegurado") or {}
    cond_a = datos.get("conductor_asegurado") or {}
    veh_t = datos.get("vehiculo_tercero") or {}
    cond_t = datos.get("conductor_tercero") or {}

    periodo = ""
    if datos.get("periodo_desde"):
        periodo = f'365 días. Desde el {datos["periodo_desde"]} hasta el {datos["periodo_hasta"]}'

    indem_a_txt = (
        "CORRESPONDE INDEMNIZAR.-" if datos.get("corresponde_indemnizar_asegurado")
        else "NO CORRESPONDE.-"
    )
    indem_t_txt = (
        "CORRESPONDE.-" if datos.get("corresponde_indemnizar_tercero")
        else "NO CORRESPONDE.-"
    )
    if datos.get("estimacion_danios_materiales"):
        indem_a_txt += f'\n  • {datos["estimacion_danios_materiales"]} para Daños Materiales'
    if datos.get("estimacion_gastos_medicos"):
        indem_a_txt += f'\n  • {datos["estimacion_gastos_medicos"]} para Gastos Médicos'
    if datos.get("exposicion_maxima_rc"):
        indem_t_txt += f'\n  • Exposición máxima RC: {datos["exposicion_maxima_rc"]}'

    filas_principales = [
        ("SOLICITANTE:", datos.get("solicitante", "")),
        ("ASEGURADO:", datos.get("asegurado", "")),
        ("N° DE PÓLIZA:", datos.get("nro_poliza", "")),
        ("SINIESTRO N°:", datos.get("nro_siniestro", "")),
        ("SECCIÓN:", datos.get("seccion", "")),
        ("PERIODO:", periodo),
        ("DETALLE DEL SINIESTRO:", datos.get("detalle_siniestro", "")),
        ("LUGAR DEL SINIESTRO:", datos.get("lugar_siniestro", "")),
        ("VEHÍCULO ASEGURADO:", _fmt_vehiculo(veh_a)),
        ("CONDUCTOR ASEGURADO:", _fmt_conductor(cond_a)),
        ("VEHÍCULO TERCERO:", _fmt_vehiculo(veh_t)),
        ("CONDUCTOR TERCERO:", _fmt_conductor(cond_t)),
        ("FECHA DEL SINIESTRO:", datos.get("fecha_siniestro", "")),
        ("FECHA DE DENUNCIA:", datos.get("fecha_denuncia", "")),
        ("FECHA DE DESIGNACIÓN:", datos.get("fecha_designacion", "")),
        ("ÚLTIMA DOCUMENTACIÓN:", datos.get("ultima_documentacion", "")),
        ("INDEMNIZACIÓN AL ASEGURADO:", indem_a_txt),
        ("INDEMNIZACIÓN AL TERCERO:", indem_t_txt),
    ]

    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in filas_principales:
        _add_data_row(tbl, label, value)
    # Anchos de columna
    for row in tbl.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(11.5)

    # ====================================================================
    # DOCUMENTACIONES COMPLEMENTARIAS
    # ====================================================================
    _heading(doc, "Documentaciones Complementarias")

    docs_comp = datos.get("documentacion_complementaria") or []
    if docs_comp:
        for d in docs_comp:
            nombre = d.get("nombre", "")
            estado = d.get("estado", "")
            try:
                p = doc.add_paragraph(style="List Bullet")
            except KeyError:
                p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.add_run(f"{nombre}: ").font.size = Pt(9.5)
            r_est = p.add_run(estado)
            r_est.bold = True
            r_est.font.size = Pt(9.5)
            r_est.font.color.rgb = VERDE if estado == "ENTREGADO" else ROJO
    else:
        doc.add_paragraph("Sin información de documentación.")

    # ====================================================================
    # ACTUACIONES
    # ====================================================================
    actuaciones = datos.get("actuaciones") or []
    if actuaciones:
        _heading(doc, "Actuaciones")
        for act in actuaciones:
            try:
                p = doc.add_paragraph(style="List Bullet")
            except KeyError:
                p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            r_f = p.add_run(f'{act.get("fecha", "")}: ')
            r_f.bold = True
            r_f.font.size = Pt(9.5)
            r_f.font.color.rgb = AZUL
            r_d = p.add_run(act.get("descripcion", ""))
            r_d.font.size = Pt(9.5)

    # ====================================================================
    # ALERTAS
    # ====================================================================
    alertas = datos.get("alertas_liquidador") or []
    if alertas:
        _heading(doc, "Alertas para el Liquidador")
        for alerta in alertas:
            _bullet(doc, alerta, color=NARANJA)

    # ====================================================================
    # SALTO DE PÁGINA — CONTRATO DE SEGURO
    # ====================================================================
    doc.add_page_break()

    _heading(doc, "Contrato de Seguro — Coberturas Afectadas")

    coberturas = datos.get("coberturas_afectadas") or []
    for cob in coberturas:
        cap = cob.get("capitulo", "")
        tit = cob.get("titulo", "")
        monto = cob.get("monto_usd", "")
        items = cob.get("items") or []

        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(6)
        r_cap = p_cap.add_run(f'Capítulo "{cap}" — {tit}')
        r_cap.bold = True
        r_cap.font.size = Pt(10)
        r_cap.font.color.rgb = AZUL

        for item in items:
            suffix = f" (hasta {monto})" if monto else ""
            _bullet(doc, f"{item}{suffix}")

    # Franquicia y exposición máxima
    p_franq = doc.add_paragraph()
    p_franq.paragraph_format.space_before = Pt(8)
    r_fl = p_franq.add_run("Franquicia: ")
    r_fl.bold = True
    p_franq.add_run(datos.get("franquicia", "—"))

    expo = datos.get("exposicion_maxima_total")
    if expo:
        p_exp = doc.add_paragraph()
        p_exp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_exp.paragraph_format.space_before = Pt(10)
        r_exp = p_exp.add_run(
            f"Siendo la exposición máxima para las coberturas afectadas, la suma de {expo}"
        )
        r_exp.bold = True
        r_exp.font.size = Pt(11)
        r_exp.font.color.rgb = ROJO

    # ====================================================================
    # SALTO DE PÁGINA — DINÁMICA DEL SINIESTRO
    # ====================================================================
    dinamica = datos.get("dinamica_siniestro", "")
    if dinamica:
        doc.add_page_break()
        _heading(doc, "Dinámica del Siniestro")
        for linea in dinamica.split("\n"):
            linea = linea.strip()
            if linea:
                p = doc.add_paragraph(linea)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(6)

    # ====================================================================
    # SALTO DE PÁGINA — CONCLUSIÓN PERICIAL
    # ====================================================================
    conclusion = datos.get("conclusion_pericial", "")
    if conclusion:
        doc.add_page_break()
        _heading(doc, "Conclusión Pericial")
        for linea in conclusion.split("\n"):
            linea = linea.strip()
            if linea:
                p = doc.add_paragraph(linea)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(6)

    resp_conclusion = datos.get("responsabilidad_conclusion", "")
    if resp_conclusion:
        p_resp = doc.add_paragraph()
        p_resp.paragraph_format.space_before = Pt(8)
        r_lbl = p_resp.add_run("Conclusión sobre responsabilidad: ")
        r_lbl.bold = True
        p_resp.add_run(resp_conclusion)

    # ====================================================================
    # INFRACCIONES Y NORMAS LEGALES
    # ====================================================================
    infracciones = datos.get("infracciones_detectadas") or []
    normas = datos.get("normas_legales_afectadas") or []

    if infracciones or normas:
        _heading(doc, "Infracciones y Normas Legales Afectadas")

        for inf in infracciones:
            p_inf = doc.add_paragraph()
            p_inf.paragraph_format.space_before = Pt(4)
            r_inf = p_inf.add_run(f'[{inf.get("infractor", "")}] ')
            r_inf.bold = True
            r_inf.font.color.rgb = ROJO
            p_inf.add_run(inf.get("descripcion_infraccion", ""))
            if inf.get("articulo_ley_5016"):
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Cm(0.5)
                r2 = p2.add_run(f'→ {inf["articulo_ley_5016"]} (Ley N° 5016/14)')
                r2.italic = True
                r2.font.size = Pt(9)
            if inf.get("articulo_ordenanza"):
                p3 = doc.add_paragraph()
                p3.paragraph_format.left_indent = Cm(0.5)
                r3 = p3.add_run(f'→ {inf["articulo_ordenanza"]}')
                r3.italic = True
                r3.font.size = Pt(9)

        if normas:
            _heading(doc, "Normas Legales Referenciadas", level=2)
            for norma in normas:
                _bullet(doc, norma)

    # ====================================================================
    # RESULTADO FINAL
    # ====================================================================
    _heading(doc, "Resultado del Análisis")

    tbl2 = doc.add_table(rows=0, cols=2)
    tbl2.style = "Table Grid"
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

    porc_a = datos.get("porcentaje_asegurado")
    porc_t = datos.get("porcentaje_tercero")
    resp_str = datos.get("responsabilidad_sugerida", "")
    if porc_a is not None:
        resp_str += f" (Asegurado {porc_a}% / Tercero {porc_t}%)"

    result_rows = [
        ("RESPONSABILIDAD SUGERIDA:", resp_str),
        ("COBERTURA APLICA:", str(datos.get("cobertura_aplica", "")).upper()),
        ("FRANQUICIA APLICA:", "SÍ" if datos.get("franquicia_aplica") else "NO"),
        ("MONTO SUGERIDO A LIQUIDAR:", datos.get("monto_sugerido_liquidar") or "—"),
        ("CONFIANZA DEL DICTAMEN:", str(datos.get("confianza_dictamen", ""))),
    ]
    for label, value in result_rows:
        _add_data_row(tbl2, label, value)
    for row in tbl2.rows:
        row.cells[0].width = Cm(6)
        row.cells[1].width = Cm(11)

    # Conflictos
    conflictos = datos.get("conflictos") or []
    if conflictos:
        _heading(doc, "Conflictos Detectados Entre Documentos", level=2)
        for c in conflictos:
            p_c = doc.add_paragraph()
            r_lbl_c = p_c.add_run(f'{c.get("campo", "")}: ')
            r_lbl_c.bold = True
            r_lbl_c.font.color.rgb = ROJO
            p_c.add_run(c.get("descripcion", ""))

    # ====================================================================
    # NOTA FINAL
    # ====================================================================
    doc.add_paragraph()
    p_nota = doc.add_paragraph()
    p_nota.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_nota.paragraph_format.space_before = Pt(10)
    r_nota = p_nota.add_run(
        "El presente Pre-Informe ha sido generado de forma automática por el sistema de inteligencia "
        "artificial de liquidación de siniestros. No es vinculante para las partes. Su contenido es una "
        "sugerencia sujeta a la aprobación y reconocimiento final de la aseguradora solicitante, quien se "
        "reserva todos los derechos y obligaciones relacionados a este siniestro y al contrato de seguros; "
        "y sin solidaridad entre sí."
    )
    r_nota.font.size = Pt(8)
    r_nota.italic = True
    r_nota.font.color.rgb = GRIS

    # ====================================================================
    # BLOQUE DE FIRMA
    # ====================================================================
    doc.add_paragraph()
    tbl_firma = doc.add_table(rows=2, cols=2)
    tbl_firma.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Espacio para la firma
    for cell in tbl_firma.rows[0].cells:
        p_sp = cell.paragraphs[0]
        p_sp.add_run("\n\n\n")
        _cell_border_bottom(cell)

    # Etiquetas
    cell_izq = tbl_firma.rows[1].cells[0]
    p_izq = cell_izq.paragraphs[0]
    p_izq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_izq = p_izq.add_run("LIQUIDADOR DE SINIESTROS\nSistema IA")
    r_izq.bold = True
    r_izq.font.size = Pt(9)

    cell_der = tbl_firma.rows[1].cells[1]
    p_der = cell_der.paragraphs[0]
    p_der.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_der = p_der.add_run(f'SOLICITANTE\n{datos.get("solicitante", "Aseguradora")}')
    r_der.bold = True
    r_der.font.size = Pt(9)

    # ====================================================================
    # SERIALIZAR A BYTES
    # ====================================================================
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
